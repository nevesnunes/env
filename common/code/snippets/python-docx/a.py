#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import requests
import io
from docx import Document
from docx.shared import Inches

url = 'https://upload.wikimedia.org/wikipedia/commons/thumb/f/f3/Usain_Bolt_Rio_100m_final_2016k.jpg/200px-Usain_Bolt_Rio_100m_final_2016k.jpg'
response = requests.get(url, stream=True)
image = io.BytesIO(response.content)

document = Document()
# || document.add_picture('/tmp/foo.jpg')
document.add_picture(image, width=Inches(1.25))
document.save('demo.docx')
